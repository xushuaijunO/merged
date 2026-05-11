"""Template parser: extracts structure, styles, and layout from a template docx."""

from dataclasses import dataclass, field
from typing import List, Dict, Optional
from docx import Document as DocxDocument


@dataclass
class StyleDef:
    """A single paragraph style definition extracted from the template."""
    name: str
    font_name: str = ""
    font_size_pt: float = 0.0
    bold: bool = False
    color: str = ""  # hex RGB
    line_spacing: float = 0.0


@dataclass
class SectionNode:
    """A section heading in the template's document structure."""
    heading: str
    level: int
    style_name: str
    children: List["SectionNode"] = field(default_factory=list)


@dataclass
class PageLayout:
    """Page dimensions and margins from the template."""
    page_width: int = 0      # EMU
    page_height: int = 0     # EMU
    margin_top: int = 0      # EMU
    margin_bottom: int = 0   # EMU
    margin_left: int = 0     # EMU
    margin_right: int = 0    # EMU


@dataclass
class TemplateSkeleton:
    """Complete template structure extracted from a docx file."""
    styles: Dict[str, StyleDef] = field(default_factory=dict)
    page_layout: PageLayout = field(default_factory=PageLayout)
    sections: List[SectionNode] = field(default_factory=list)
    cover_elements: List[dict] = field(default_factory=list)
    has_header: bool = False
    has_footer: bool = False

    def get_body_sections(self) -> List[SectionNode]:
        """Return main body chapter headings.

        Body chapters are Heading 2 style entries. Excludes:
        - 一级条标题 entries (these are sub-sections under a parent chapter)
        - Appendix / cover / TOC entries
        """
        for section in self.sections:
            if section.children and any(
                c.style_name == "Heading 2" for c in section.children
            ):
                # Return only Heading 2 children (true chapter headings)
                return [c for c in section.children if c.style_name == "Heading 2"]

        # Fallback: return second-level headings excluding 一级条标题
        result = []
        for section in self.sections:
            for child in section.children:
                if (child.level >= 2 and
                    not child.heading.startswith("附件") and
                    child.style_name != "一级条标题"):
                    result.append(child)
        return result


STYLE_LEVEL_MAP = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "章标题": 1, "一级条标题": 2, "二级条标题": 3, "三级条标题": 4,
    "前言、引言标题": 1, "封面标准名称": 1,
}


def _get_level(style_name: str) -> int:
    return STYLE_LEVEL_MAP.get(style_name, 0)


def parse_template(filepath: str) -> TemplateSkeleton:
    """Parse a template docx and extract its full skeleton."""
    doc = DocxDocument(filepath)
    skel = TemplateSkeleton()

    # --- Extract style definitions ---
    for s in doc.styles:
        if s.type is None or 'PARAGRAPH' not in str(s.type):
            continue
        try:
            font = s.font
            sd = StyleDef(name=s.name)
            if font.name:
                sd.font_name = font.name
            if font.size:
                sd.font_size_pt = font.size.pt
            if font.bold is not None:
                sd.bold = font.bold
            if font.color and font.color.rgb:
                sd.color = str(font.color.rgb)
            pf = s.paragraph_format
            if pf and pf.line_spacing:
                sd.line_spacing = float(pf.line_spacing)
            if sd.font_name or sd.font_size_pt:
                skel.styles[s.name] = sd
        except Exception:
            pass

    # --- Extract page layout ---
    if doc.sections:
        sec = doc.sections[0]
        skel.page_layout = PageLayout(
            page_width=sec.page_width,
            page_height=sec.page_height,
            margin_top=sec.top_margin,
            margin_bottom=sec.bottom_margin,
            margin_left=sec.left_margin,
            margin_right=sec.right_margin,
        )
        # Check header/footer
        try:
            if sec.header and any(p.text.strip() for p in sec.header.paragraphs):
                skel.has_header = True
        except Exception:
            pass
        try:
            if sec.footer and any(p.text.strip() for p in sec.footer.paragraphs):
                skel.has_footer = True
        except Exception:
            pass

    # --- Extract section tree ---
    section_stack: List[SectionNode] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style else ""
        level = _get_level(style_name)
        if level > 0:
            node = SectionNode(heading=text, level=level, style_name=style_name)
            while section_stack and section_stack[-1].level >= level:
                section_stack.pop()
            if section_stack:
                section_stack[-1].children.append(node)
            else:
                skel.sections.append(node)
            section_stack.append(node)

    # --- Extract cover elements (paragraphs before first real heading) ---
    found_first_heading = False
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""
        level = _get_level(style_name)
        if level > 0 and style_name not in ("封面标准名称", "封面标准号2",
                                              "其他标准标志", "其他标准称谓"):
            found_first_heading = True
        if found_first_heading:
            break
        if text:
            skel.cover_elements.append({
                "text": text,
                "style_name": style_name,
            })

    return skel
