"""Document parser: extracts structured content from docx files.

Enhanced with perceptual hashing (dHash) for image deduplication and
precise paragraph-level positioning for smart image insertion.
"""

import io
import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple
from docx import Document as DocxDocument
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Perceptual hashing (dHash) — implemented with Pillow only, no imagehash lib
# ---------------------------------------------------------------------------

def compute_dhash(image_blob: bytes, hash_size: int = 8) -> str:
    """Compute difference hash for an image blob.

    Returns a hex string. Images with Hamming distance < 5 are considered
    visually identical or near-identical (same image with minor compression
    or watermark differences).
    """
    from PIL import Image as PILImage

    img = PILImage.open(io.BytesIO(image_blob))
    # Convert to grayscale, resize to hash_size+1 x hash_size
    img = img.convert("L").resize((hash_size + 1, hash_size), PILImage.LANCZOS)
    pixels = list(img.getdata())

    # Compute difference: compare adjacent pixels
    diff = 0
    for row in range(hash_size):
        for col in range(hash_size):
            idx = row * (hash_size + 1) + col
            diff <<= 1
            if pixels[idx] > pixels[idx + 1]:
                diff |= 1

    return hex(diff)[2:].zfill(hash_size * hash_size // 4)


def hamming_distance(h1: str, h2: str) -> int:
    """Compute Hamming distance between two hex hash strings."""
    if len(h1) != len(h2):
        return 999
    dist = 0
    for c1, c2 in zip(h1.lower(), h2.lower()):
        xor = int(c1, 16) ^ int(c2, 16)
        dist += bin(xor).count("1")
    return dist


def deduplicate_images(images: List["Image"], threshold: int = 5) -> List[List["Image"]]:
    """Group images by perceptual similarity.

    Returns a list of groups, each group contains visually similar images.
    The first image in each group is the "canonical" one (largest file).
    """
    if not images:
        return []

    # Group by dHash similarity
    groups = []
    remaining = list(images)

    while remaining:
        pivot = remaining.pop(0)
        group = [pivot]
        new_remaining = []
        for img in remaining:
            if hamming_distance(pivot.dhash, img.dhash) <= threshold:
                group.append(img)
            else:
                new_remaining.append(img)
        remaining = new_remaining
        # Sort: largest (best quality) first
        group.sort(key=lambda x: len(x.blob), reverse=True)
        groups.append(group)

    return groups


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class Image:
    """Represents an extracted image with precise positioning context."""
    blob: bytes
    content_type: str
    filename: str
    dhash: str = ""                          # Perceptual hash for dedup
    section_heading: str = ""                # Heading of the parent section
    context_before: str = ""                 # Paragraph text immediately before image
    context_after: str = ""                  # Paragraph text immediately after image
    paragraph_index: int = -1                # Position in source document
    source_doc: str = ""                     # Which document this image came from
    caption: str = ""                        # Best-guess image caption (auto-extracted)


@dataclass
class Table:
    """Represents a table."""
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class Section:
    """A document section with hierarchical structure."""
    heading: str
    level: int
    paragraphs: List[str] = field(default_factory=list)
    tables: List[Table] = field(default_factory=list)
    images: List[Image] = field(default_factory=list)
    children: List["Section"] = field(default_factory=list)
    # Ordered body: list of dicts {"type": "text"|"image", "value": str|Image}
    # Lets the renderer interleave paragraphs and images as they appear in source.
    body: List[Dict[str, Any]] = field(default_factory=list)
    # Original docx style name for this section's heading — preserved so the
    # merger can filter cover/preface/TOC sections by their authoring style.
    style_name: str = ""


@dataclass
class ParsedDocument:
    """Complete parsed document with images."""
    filename: str
    title: str
    full_text: str
    source_path: str = ""
    sections: List[Section] = field(default_factory=list)
    all_images: List[Image] = field(default_factory=list)
    all_tables: List[Table] = field(default_factory=list)

    def to_dict(self) -> dict:
        """Serialize sections for the analyzer and merger."""
        def section_to_dict(sec: Section) -> dict:
            return {
                "heading": sec.heading,
                "level": sec.level,
                "style_name": sec.style_name,
                "paragraphs": sec.paragraphs,
                "tables": [t.rows for t in sec.tables],
                "image_count": len(sec.images),
                "images": [
                    {
                        "dhash": img.dhash,
                        "context_before": img.context_before,
                        "context_after": img.context_after,
                        "section_heading": img.section_heading,
                        "source_doc": img.source_doc,
                        "content_type": img.content_type,
                        "filename": img.filename,
                        "caption": img.caption,
                        "size_bytes": len(img.blob),
                    }
                    for img in sec.images
                ],
                "body": [
                    # Serialize body items keeping references to image objects
                    # so the merger can re-render them in source order.
                    {"type": item["type"], "value": item["value"]}
                    for item in sec.body
                ],
                "children": [section_to_dict(c) for c in sec.children],
            }

        def collect_section_images(sec: Section, result: dict):
            if sec.images:
                result[sec.heading] = sec.images
            for child in sec.children:
                collect_section_images(child, result)

        section_images = {}
        for s in self.sections:
            collect_section_images(s, section_images)

        return {
            "filename": self.filename,
            "title": self.title,
            "source_path": self.source_path,
            "sections": [section_to_dict(s) for s in self.sections],
            "table_count": len(self.all_tables),
            "image_count": len(self.all_images),
            "section_images": section_images,
        }


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------

HEADING_STYLE_MAP = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "章标题": 1,
    "一级条标题": 2,
    "二级条标题": 3,
    "三级条标题": 4,
    "四级条标题": 5,
    "前言、引言标题": 1,
    "封面标准名称": 1,
    "附录标题": 1,
}


def get_heading_level(style_name: str) -> int:
    """Convert a Word style name to a heading level number."""
    if style_name in HEADING_STYLE_MAP:
        return HEADING_STYLE_MAP[style_name]
    if style_name.startswith("Heading") or style_name.startswith("heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            pass
    if "标题" in style_name or "Heading" in style_name:
        return 2
    return 0


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def extract_images(doc: DocxDocument,
                   source_doc: str = "") -> Tuple[List[Image], Dict[str, Image]]:
    """Extract all images from a docx document with dHash computation."""
    images = []
    rid_map = {}
    for rId, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_part = rel.target_part
            blob = image_part.blob
            dhash = compute_dhash(blob)
            img = Image(
                blob=blob,
                content_type=image_part.content_type,
                filename=image_part.partname.split("/")[-1],
                dhash=dhash,
                source_doc=source_doc,
            )
            images.append(img)
            rid_map[rId] = img
    return images, rid_map


def _extract_inline_image_rid(paragraph) -> Optional[str]:
    """Check if a paragraph contains an inline image and return its rId."""
    for blip in paragraph._element.findall('.//' + qn('a:blip')):
        embed = blip.get(qn('r:embed'))
        if embed:
            return embed
    return None


# Regex for detecting explicit image captions in surrounding paragraphs.
# Matches: "图 1 xxx", "图1.2 xxx", "图 2-3 xxx", "Figure 1 xxx", etc.
_CAPTION_PATTERNS = [
    re.compile(r'^图\s*[\d\.\-]+\s*[:：]?\s*(.+)$'),
    re.compile(r'^Figure\s+[\d\.\-]+\s*[:：]?\s*(.+)$', re.IGNORECASE),
    re.compile(r'^Fig\.?\s*[\d\.\-]+\s*[:：]?\s*(.+)$', re.IGNORECASE),
]


def _detect_caption(context_before: str, context_after: str) -> str:
    """Return the caption text if a 图N/Figure N pattern is found nearby.

    Empty string means no explicit caption detected — caller should fall back
    to summarized context.
    """
    for candidate in (context_after, context_before):
        if not candidate:
            continue
        text = candidate.strip()
        for pat in _CAPTION_PATTERNS:
            m = pat.match(text)
            if m:
                # Return just the descriptive part (after "图N")
                tail = m.group(1).strip()
                return tail or text
    return ""


def _summarize_caption(section_heading: str, context_before: str,
                        context_after: str) -> str:
    """Build a short caption from context when no explicit 图N label exists."""
    for candidate in (context_after, context_before):
        text = (candidate or "").strip()
        if not text:
            continue
        # Take the first sentence-ish chunk, max 30 chars
        for sep in ("。", "；", ";", "\n", "，", ","):
            if sep in text:
                text = text.split(sep, 1)[0]
                break
        text = text.strip()
        if len(text) > 30:
            text = text[:30] + "…"
        if text:
            return text
    if section_heading:
        return section_heading
    return "示意图"


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_document(filepath: str, filename: str) -> ParsedDocument:
    """Parse a docx file into structured representation with precise image tracking."""
    doc = DocxDocument(filepath)

    # Extract images with dHash
    images, rid_map = extract_images(doc, source_doc=filename)

    # Walk doc.element.body to get paragraphs and tables in source order
    P_TAG = qn('w:p')
    TBL_TAG = qn('w:tbl')

    body_iter = list(doc.element.body.iterchildren())
    para_pool = iter(doc.paragraphs)
    tbl_pool = iter(doc.tables)

    body_items = []
    for child in body_iter:
        if child.tag == P_TAG:
            para = next(para_pool, None)
            if para is None:
                continue
            text = para.text.strip()
            img_rid = _extract_inline_image_rid(para)
            body_items.append({
                "kind": "para",
                "text": text,
                "has_image": img_rid is not None and img_rid in rid_map,
                "img_rid": img_rid,
                "style_name": para.style.name if para.style else "",
            })
        elif child.tag == TBL_TAG:
            tbl = next(tbl_pool, None)
            if tbl is None:
                continue
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            body_items.append({"kind": "tbl", "rows": rows})

    # Legacy: flat list of all paragraphs (for caption pre-scan and context lookup)
    all_paragraphs = [item for item in body_items if item["kind"] == "para"]

    # Pre-scan: identify caption paragraphs (图N pattern adjacent to an image).
    # These should not appear as regular body text — they're consumed as image captions.
    caption_indices = set()
    for i, pdata in enumerate(all_paragraphs):
        if pdata["has_image"] or not pdata["text"]:
            continue
        is_caption_format = any(
            pat.match(pdata["text"]) for pat in _CAPTION_PATTERNS
        )
        if not is_caption_format:
            continue
        # Caption must be adjacent (within 1 paragraph either side) to an image
        for j in (i - 1, i + 1):
            if 0 <= j < len(all_paragraphs) and all_paragraphs[j]["has_image"]:
                caption_indices.add(i)
                break

    # Build document-level tables list (for backward compat with ParsedDocument.all_tables)
    tables = [Table(rows=item["rows"]) for item in body_items if item["kind"] == "tbl"]

    # Build section tree by walking the interleaved body
    root_sections = []
    section_stack: List[Section] = []
    title = filename.replace(".docx", "").strip()
    para_idx = -1  # Tracks legacy paragraph index (for image positioning)

    for body_item in body_items:
        # ---- TABLE ----
        if body_item["kind"] == "tbl":
            current_section = section_stack[-1] if section_stack else None
            tbl_obj = Table(rows=body_item["rows"])
            if current_section:
                current_section.tables.append(tbl_obj)
                current_section.body.append({"type": "table", "value": tbl_obj})
            else:
                # Pre-heading table — attach to preamble
                if root_sections and root_sections[-1].heading == "_preamble":
                    root_sections[-1].tables.append(tbl_obj)
                    root_sections[-1].body.append({"type": "table", "value": tbl_obj})
                else:
                    preamble = Section(heading="_preamble", level=0)
                    preamble.tables.append(tbl_obj)
                    preamble.body.append({"type": "table", "value": tbl_obj})
                    root_sections.append(preamble)
            continue

        # ---- PARAGRAPH ----
        para_idx += 1
        pdata = body_item
        text = pdata["text"]
        style_name = pdata["style_name"]
        has_image = pdata["has_image"]
        img_rid = pdata["img_rid"]
        level = get_heading_level(style_name)

        # --- Get context: text of surrounding paragraphs ---
        context_before = ""
        context_after = ""

        if has_image:
            # Look backward for the nearest non-empty text
            for j in range(para_idx - 1, -1, -1):
                bt = all_paragraphs[j]["text"]
                if bt:
                    context_before = bt
                    break
            # Look forward for the nearest non-empty text
            for j in range(para_idx + 1, len(all_paragraphs)):
                at = all_paragraphs[j]["text"]
                if at:
                    context_after = at
                    break

        if not text and not has_image:
            continue

        if level > 0:
            new_section = Section(heading=text, level=level, style_name=style_name)

            while section_stack and section_stack[-1].level >= level:
                section_stack.pop()

            if section_stack:
                section_stack[-1].children.append(new_section)
            else:
                root_sections.append(new_section)

            section_stack.append(new_section)

            if style_name == "封面标准名称":
                title = text
        else:
            current_section = section_stack[-1] if section_stack else None
            current_heading = current_section.heading if current_section else ""

            # --- Handle image paragraph ---
            if has_image:
                img = rid_map[img_rid]
                img.section_heading = current_heading
                img.context_before = context_before
                img.context_after = context_after
                img.paragraph_index = para_idx

                # Also record the heading as context if no paragraph-level context
                if not img.context_before and current_section:
                    img.context_before = current_section.heading

                # Caption detection: explicit 图N pattern first, else summarize
                explicit = _detect_caption(context_before, context_after)
                if explicit:
                    img.caption = explicit
                else:
                    img.caption = _summarize_caption(
                        current_heading, context_before, context_after,
                    )

                if current_section:
                    current_section.images.append(img)
                    current_section.body.append({"type": "image", "value": img})
                else:
                    preamble = Section(heading="_preamble", level=0)
                    preamble.images.append(img)
                    preamble.body.append({"type": "image", "value": img})
                    root_sections.append(preamble)

            # --- Handle text paragraph ---
            if text and para_idx not in caption_indices:
                if current_section:
                    current_section.paragraphs.append(text)
                    current_section.body.append({"type": "text", "value": text})
                else:
                    if root_sections and root_sections[-1].heading == "_preamble":
                        root_sections[-1].paragraphs.append(text)
                        root_sections[-1].body.append({"type": "text", "value": text})
                    else:
                        preamble = Section(heading="_preamble", level=0)
                        preamble.paragraphs.append(text)
                        preamble.body.append({"type": "text", "value": text})
                        root_sections.append(preamble)

    # Build full text
    full_text_parts = []
    for sec in root_sections:
        full_text_parts.append(sec.heading)
        full_text_parts.extend(sec.paragraphs)
    full_text = "\n".join(full_text_parts)

    return ParsedDocument(
        filename=filename,
        title=title,
        full_text=full_text,
        source_path=filepath,
        sections=root_sections,
        all_images=images,
        all_tables=tables,
    )
