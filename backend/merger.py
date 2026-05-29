"""Merger: generates the final unified operating procedure docx.

Heading hierarchy (all black, visible in Word navigation pane):
  Heading 1: 目录, 前言, 1-7 body chapters, 附件A-N titles
  Heading 2: sub-sections (1.1, 1.2...), appendix top-level sections
  Heading 3: sub-sub-sections (1.1.1...)
"""

import os
import io as std_io
import re
import shutil
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from analyzer import MergePlan

BLACK = RGBColor(0, 0, 0)


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting markers: **bold**, *italic*, __bold__, _italic_."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    return text.strip()


# ====================================================================
# Heading helpers — all black, all use Word heading styles for nav pane
# ====================================================================

def _h1(doc, text):
    """Heading 1: 目录, 前言, chapter titles, appendix titles."""
    h = doc.add_heading(_strip_markdown(text), level=1)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = "黑体"
    return h


def _h2(doc, text):
    """Heading 2: sub-sections like 1.1, 1.2."""
    h = doc.add_heading(_strip_markdown(text), level=2)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = "黑体"
    return h


def _h3(doc, text):
    """Heading 3: sub-sub-sections like 1.1.1."""
    h = doc.add_heading(_strip_markdown(text), level=3)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = "黑体"
    return h


def _h4(doc, text):
    """Heading 4: deepest sub-sections like 1.1.1.1 (used in appendix only)."""
    h = doc.add_heading(_strip_markdown(text), level=4)
    for run in h.runs:
        run.font.color.rgb = BLACK
        run.font.name = "黑体"
    return h


def _body(doc, text):
    """Normal body paragraph (宋体 10.5pt).

    Automatically strips leading heading-numbering patterns (e.g. "1.1 内容")
    from body paragraphs. Preserves **bold** markers as actual bold runs.
    """
    # Strip "N.M content" / "N.M.K content" etc. from the start of body text
    cleaned = re.sub(r'^\d+(?:\.\d+)+\s+', '', text).strip()
    p = doc.add_paragraph()
    # Split on **bold** markers — odd-indexed parts are bold
    parts = re.split(r'\*\*(.+?)\*\*', cleaned)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run.font.color.rgb = BLACK
        if i % 2 == 1:
            run.bold = True
    return p


def _table(doc, rows: List[List[str]]):
    """Insert a table with header row styled."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < ncols:
                cell = tbl.rows[ri].cells[ci]
                cell.text = str(cell_text)
                if ri == 0:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(9)
    doc.add_paragraph()


def _add_image_to_doc(doc, image_blob, content_type, caption="", width_inches=5.0):
    image_stream = std_io.BytesIO(image_blob)
    try:
        doc.add_picture(image_stream, width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.name = "宋体"
            run.font.color.rgb = BLACK
            run.italic = True
    except Exception:
        pass


# ====================================================================
# Template-as-base cover preservation
# ====================================================================

# Boundary signals: paragraph styles that mark the END of the cover area
_POST_COVER_STYLES = {
    'Heading 1', 'Heading 2', '章标题', '一级条标题',
    '前言、引言标题', '目录、目次标题',
    'toc 1', 'toc 2', 'toc 3', 'TOC 1', 'TOC 2', 'TOC 3',
}


def _trim_after_cover(doc):
    """Remove all body content after the cover area, keeping sectPr intact.

    Walks the docx body in order. Stops at the first paragraph that signals
    post-cover content:
    - Paragraph style matching a known TOC/preface/body style
    - Paragraph text containing '目录' / '目次' / '前言' (whitespace-tolerant)
    - A TOC field code (instrText with TOC)

    Everything from that point up to (but not including) sectPr is deleted.
    Covers may include tables and section breaks — we skip past those.
    """
    body = doc.element.body
    children = list(body.iterchildren())

    cover_end_idx = None
    for i, child in enumerate(children):
        if i > 60:
            break
        tag = child.tag

        # Skip tables — covers often have decorative tables
        if qn('w:tbl') in tag:
            continue

        # Skip section breaks — cover may span multiple sections
        if qn('w:sectPr') in tag:
            continue

        # Only process paragraphs
        if qn('w:p') not in tag:
            continue

        text = ''.join(child.itertext())
        text_norm = re.sub(r'\s+', '', text)

        pStyle = child.find('.//' + qn('w:pStyle'))
        style_name = pStyle.get(qn('w:val'), '') if pStyle is not None else ''

        # Check style-based boundary
        if style_name in _POST_COVER_STYLES:
            cover_end_idx = i
            break

        # Check for TOC field code
        instr = child.find('.//' + qn('w:instrText'))
        if instr is not None and instr.text and 'TOC' in instr.text:
            cover_end_idx = i
            break

        # Check whitespace-normalized keyword matching
        if text_norm in ('目录', '目次', '前言', '引言'):
            cover_end_idx = i
            break

        # Short paragraph with keyword (handles "目    录" style variations)
        if text_norm and len(text_norm) < 10 and (
            '目录' in text_norm or '目次' in text_norm or '前言' in text_norm
        ):
            cover_end_idx = i
            break

    if cover_end_idx is None:
        return  # No boundary detected — leave template body intact

    # Delete from cover_end_idx onwards, except sectPr
    for child in children[cover_end_idx:]:
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def _replace_cover_title(doc, new_title):
    """Replace the cover title text while preserving original run formatting."""
    if not new_title:
        return
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "封面标准名称":
            if para.runs:
                first_run = para.runs[0]
                first_run.text = new_title
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_title)
            return


_STYLE_FONTS = {
    'Heading 1': (1, '黑体', 16),
    'Heading 2': (2, '黑体', 14),
    'Heading 3': (3, '黑体', 12),
    'Heading 4': (4, '黑体', 12),
}


def _detach_heading_numbering(doc):
    """Remove multilevel numbering from all heading styles.

    The template's heading styles (Heading 1/2/3/4) may have w:numPr defined
    in their style XML, linking them to multilevel list numbering. This causes
    Word to ADD auto-numbering on top of our text-based numbering, creating
    the dreaded "双重编号" (e.g. Word shows "1 1.1 岗位与巡检").
    """
    from lxml import etree
    # Remove numPr from ANY style whose name indicates it's a heading/标题
    # style (matches Heading 1-9, 章标题, 一级条标题, 标题1-4, etc.)
    # Stops Word from stacking auto-numbering on top of our text numbering.
    for style in doc.styles:
        name = style.name or ""
        has_heading_keyword = (
            'Heading' in name or 'heading' in name or
            '标题' in name or '章' in name or
            '条标题' in name or
            'toc' in name.lower() or 'TOC' in name
        )
        if has_heading_keyword:
            try:
                for numpr in style.element.findall('.//' + qn('w:numPr')):
                    numpr.getparent().remove(numpr)
            except Exception:
                pass


def _ensure_styles(doc):
    """Create standard heading styles if the document lacks them.

    Template documents often use numeric or Chinese style IDs ('1', '2', '章标题')
    rather than the built-in 'Heading 1/2/3/4' names. Since downstream pipeline
    code (`_h1`, `_h2`, etc.) uses `add_heading(text, level=...)`, we must
    ensure those style names are available.
    """
    from lxml import etree
    existing_names = {s.name for s in doc.styles}

    for style_name, (lvl, font_name, font_size) in _STYLE_FONTS.items():
        if style_name in existing_names:
            continue
        try:
            s = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            s.font.name = font_name
            s.font.size = Pt(font_size)
            s.font.bold = True
            try:
                s.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            except Exception:
                pass
            # Set outlineLvl so Word nav pane / TOC pick it up
            pPr = s.element.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(s.element, qn('w:pPr'))
            outline = etree.SubElement(pPr, qn('w:outlineLvl'))
            outline.set(qn('w:val'), str(lvl - 1))
        except Exception:
            pass


# ====================================================================
# Default cover (when no template is available)
# ====================================================================

def _create_default_cover(doc, cover_title=""):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("上海浦东威立雅自来水有限公司企业标准")
    run.font.size = Pt(16)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cover_title or "操作规程")
    run.font.size = Pt(26)
    run.bold = True
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("上海浦东威立雅自来水有限公司  发布")
    run.font.size = Pt(14)
    doc.add_page_break()


# ====================================================================
# TOC — Word TOC field (auto-generates on open)
# ====================================================================

def _insert_toc_field(doc):
    """Insert a Word TOC field code so the table of contents is auto-generated.

    Uses TOC field with the `\n` switch (no hyperlinks) to avoid the
    "Error! Bookmark not defined" compatibility issue in Word 365+.
    The TOC refreshes automatically when the user opens the document in Word.
    """
    from lxml import etree
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    h1_para = doc.add_heading("目    录", level=1)
    for run in h1_para.runs:
        run.font.color.rgb = BLACK
        run.font.name = "黑体"
    doc.add_paragraph()

    p = doc.add_paragraph()

    # w:fldChar = begin
    begin_run = p.add_run()
    begin_fld = etree.SubElement(begin_run._element, qn('w:fldChar'))
    begin_fld.set(qn('w:fldCharType'), 'begin')

    # w:instrText = TOC fields
    instr_run = p.add_run()
    instr = etree.SubElement(instr_run._element, qn('w:instrText'))
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u \\n '

    # w:fldChar = separate
    sep_run = p.add_run()
    sep_fld = etree.SubElement(sep_run._element, qn('w:fldChar'))
    sep_fld.set(qn('w:fldCharType'), 'separate')

    # Placeholder text (Word replaces it on field update)
    placeholder = p.add_run("（请更新域以生成目录 / 右键 → 更新域）")
    placeholder.font.size = Pt(10)
    placeholder.font.name = "宋体"
    placeholder.font.color.rgb = RGBColor(128, 128, 128)
    placeholder.italic = True

    # w:fldChar = end
    end_run = p.add_run()
    end_fld = etree.SubElement(end_run._element, qn('w:fldChar'))
    end_fld.set(qn('w:fldCharType'), 'end')

    doc.add_page_break()


# ====================================================================
# Preface
# ====================================================================

def _generate_preface(doc, template_path=None):
    doc.add_page_break()
    _h1(doc, "前    言")
    doc.add_paragraph()

    if template_path and os.path.exists(template_path):
        src_doc = Document(template_path)
        in_preface = False
        copied = 0
        for para in src_doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if "前言" in text and ("前言、引言标题" in style or len(text) <= 6):
                in_preface = True
                continue
            if in_preface:
                if style in ("Heading 1", "Heading 2", "章标题", "一级条标题",
                              "前言、引言标题") and text and "前言" not in text:
                    break
                if text:
                    _body(doc, text)
                    copied += 1
                    if copied > 10:
                        break

    last_text = doc.paragraphs[-1].text.strip() if doc.paragraphs else ""
    if not last_text or last_text == "前    言":
        _body(doc, "本标准按照企业标准编写规范起草。")
        _body(doc, "本标准由上海浦东威立雅自来水有限公司浦东水厂生产管理科提出并归口。")
        _body(doc, "本标准起草部门：浦东水厂生产管理科。")

    doc.add_page_break()


_PREFACE_PATTERNS = ('前言', '前　言', '前    言', '引言')
_TOC_PATTERNS = ('目录', '目次')
_ws_re = re.compile(r'\s+')
_FRONT_MATTER_SKIP = {'前言', '目录', '目次', '引言', '前　言'}
_STRIP_NUM = re.compile(r'^(\d+\.)+\d+\s*').sub  # strip "6.1 "/"1.1综合" → "接收"/"综合数据查询"
_STRIP_CN_NUM = re.compile(r'^[一二三四五六七八九十]+[、．.\s]+').sub  # strip "一、概述" → "概述"
_STRIP_SIMPLE_NUM = re.compile(r'^\d+[、．.)、]\s*').sub  # strip "1、进入" → "进入"


def _is_preface_heading(heading: str) -> bool:
    """Check if a heading is preface — handles whitespace variations."""
    if not heading:
        return False
    norm = _ws_re.sub('', heading)
    return any(p.replace(' ', '').replace('　', '') in norm for p in _PREFACE_PATTERNS)


def _is_toc_heading(heading: str) -> bool:
    if not heading:
        return False
    norm = _ws_re.sub('', heading)
    return any(p in norm for p in _TOC_PATTERNS)


def _is_front_matter_heading(heading: str) -> bool:
    return _is_preface_heading(heading) or _is_toc_heading(heading)


# ====================================================================
# Appendix content cleaning
# ====================================================================


def _norm(s: str) -> str:
    """Normalize a heading for comparison: strip whitespace and leading numbering."""
    s = _ws_re.sub('', s).strip()
    s = re.sub(r'^[\d\.、]+', '', s).strip()
    return s


# Universal boilerplate chapters — filtered from appendices because they
# duplicate main body content. Operational chapters (作业要求, 应急处置, etc.)
# are kept so detailed steps survive.
_UNIVERSAL_BOILERPLATE = {
    '前言', '目录', '目次', '引言',
    '范围', '规范性引用文件', '规范性引用文献',
    '术语和定义', '术语', '定义',
    '缩略语', '符号和缩略语', '符号',
}

# Authoring styles that mark front matter in standard docx templates.
# Any section whose original style name is one of these gets dropped entirely
# from the appendix, even if its heading text doesn't match a known keyword.
_FRONT_MATTER_STYLES = {
    '封面标准名称', '封面标准号', '封面标准号2',
    '其他标准标志', '其他标准称谓',
    '其他发布日期', '其他实施日期',
    '前言、引言标题',
    '目录、目次标题',
    'toc 1', 'toc 2', 'toc 3',
    'TOC 1', 'TOC 2', 'TOC 3',
}

# Regex matching the first "real content" chapter heading.
# Examples: "1 范围", "1.范围", "1、范围", "1） 范围", "一 范围", "第一章 ..."
_NUMBERED_CHAPTER_RE = re.compile(
    r'^\s*(?:第\s*[一二三四五六七八九十0-9]+\s*(?:章|节|条)|'
    r'[1-9一二三四五六七八九][\d]*)\s*[\.\、\)\）\s]'
)


def _is_numbered_chapter(heading: str) -> bool:
    if not heading:
        return False
    return bool(_NUMBERED_CHAPTER_RE.match(heading.strip()))


def _should_skip_section(heading: str, body_headings: set, style_name: str = "") -> bool:
    """Check if a section heading should be skipped in appendix rendering.

    Skips: empty/preamble headings, front matter (preface/TOC), universal
    boilerplate chapters, and any section whose original docx style marks it
    as cover/preface/TOC content.
    """
    if style_name and style_name in _FRONT_MATTER_STYLES:
        return True
    h_norm = _norm(heading)
    if not h_norm or h_norm == '_preamble':
        return True
    if h_norm in _FRONT_MATTER_SKIP:
        return True
    if _is_front_matter_heading(heading):
        return True
    if h_norm in _UNIVERSAL_BOILERPLATE:
        return True
    if '见表' in h_norm:
        return True
    if re.match(r'^表\s*\d+', h_norm):
        return True
    if h_norm.startswith('本文件规定') or h_norm.startswith('下列文件'):
        return True
    if re.match(r'^[A-Z]+-[A-Z]-[A-Z]?\d+', h_norm):
        return True
    return False


def _find_content_start_index(sections: List[dict]) -> int:
    """Find the index of the first 'real content' section.

    Scans top-level sections for the first one whose heading starts with a
    chapter number (1 / 一 / 第一章 / etc). Everything before that is treated
    as front matter (cover + preface + TOC + bookkeeping) and dropped.

    Returns 0 if no numbered chapter is found — caller falls back to per-section
    filtering only.
    """
    for i, sec in enumerate(sections):
        heading = (sec.get("heading", "") or "").strip()
        if _is_numbered_chapter(heading):
            return i
    return 0


def _render_appendix_sections(doc, sections: List[dict], body_headings: set,
                               image_map: dict, counters=None,
                               apply_content_boundary: bool = True):
    """Recursively render source doc sections as appendix content.

    Uses section-level-based counting: each section's heading depth is
    determined by its own ``level`` from parsing, not by tree position.
    This ensures consistent numbering regardless of which parent sections
    were skipped as front matter or boilerplate.

    Counter mapping (base level = 2):
      level 2 → counter[0] → H2 "N"
      level 3 → counter[1] → H3 "N.M"
      level 4 → counter[2] → H4 "N.M.K"
      level 5+ → bold body text (no heading number)
    """
    if counters is None:
        counters = [0]

    # Top-level boundary: drop everything before the first "1 xxx" / "一 xxx"
    if apply_content_boundary:
        start = _find_content_start_index(sections)
        if start > 0:
            sections = sections[start:]

    for sec in sections:
        sec_heading = sec.get("heading", "").strip()
        sec_style = sec.get("style_name", "") or ""
        section_level = sec.get("level", 0)

        if _should_skip_section(sec_heading, body_headings, sec_style):
            # Section heading is skipped. For table-related skips, preserve
            # body content (tables, text) and consume a counter so children
            # maintain correct numbering. For universal boilerplate (范围,
            # 规范性引用文件, etc.), skip entirely — no counter, no body.
            h_norm = _norm(sec_heading)
            is_boilerplate = (h_norm in _UNIVERSAL_BOILERPLATE or
                             _is_front_matter_heading(sec_heading) or
                             h_norm == '_preamble' or
                             sec_style in _FRONT_MATTER_STYLES or
                             h_norm.startswith('本文件规定') or
                             h_norm.startswith('下列文件') or
                             bool(re.match(r'^[A-Z]+-[A-Z]-[A-Z]?\d+', h_norm)))
            body_items = sec.get("body", [])
            children = sec.get("children", [])
            if body_items and not is_boilerplate:
                for item in body_items:
                    kind = item.get("type")
                    value = item.get("value")
                    if kind == "text" and value and str(value).strip():
                        _body(doc, str(value).strip())
                    elif kind == "image" and value is not None:
                        blob = getattr(value, "blob", None)
                        if blob:
                            _add_image_to_doc(doc, blob, getattr(value, "content_type", ""),
                                              caption=getattr(value, "caption", "") or "")
                    elif kind == "table" and value is not None:
                        rows = getattr(value, "rows", None) or []
                        if rows:
                            _table(doc, rows)
            elif not is_boilerplate:
                for p_text in sec.get("paragraphs", []):
                    if p_text.strip():
                        _body(doc, p_text.strip())
                for tbl_data in sec.get("tables", []):
                    if isinstance(tbl_data, list) and tbl_data:
                        _table(doc, tbl_data)
            if children:
                _render_appendix_sections(
                    doc, children, body_headings, image_map, counters,
                    apply_content_boundary=False,
                )
            continue

        has_heading = bool(sec_heading)
        if has_heading:
            # Use section's own level for counter depth, not tree position
            counter_idx = section_level - 2  # L2→0, L3→1, L4→2
            # Strip source-document numbering from heading text
            clean_sec_title = _STRIP_NUM('', sec_heading)
            clean_sec_title = _STRIP_CN_NUM('', clean_sec_title)
            clean_sec_title = _STRIP_SIMPLE_NUM('', clean_sec_title)
            clean_sec_title = clean_sec_title.strip()
            if not clean_sec_title:
                clean_sec_title = sec_heading.strip()

            if counter_idx < 0 or counter_idx >= 3:
                # Preamble/cover (L<=1) or deep sub-steps (L>=5):
                # render as bold body text, no heading number
                p = doc.add_paragraph()
                run = p.add_run(clean_sec_title)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = "宋体"
                run.font.color.rgb = BLACK
            else:
                # Ensure counter array is wide enough
                while len(counters) <= counter_idx:
                    counters.append(0)
                counters[counter_idx] += 1
                # Reset deeper counters
                for i in range(counter_idx + 1, len(counters)):
                    counters[i] = 0
                num_str = ".".join(str(c) for c in counters[:counter_idx + 1])
                if counter_idx == 0:
                    _h2(doc, f"{num_str} {clean_sec_title}")
                elif counter_idx == 1:
                    _h3(doc, f"{num_str} {clean_sec_title}")
                else:  # counter_idx == 2
                    _h4(doc, f"{num_str} {clean_sec_title}")

        # Prefer ordered body (interleaved text/image/table), fall back to legacy fields
        body_items = sec.get("body", [])
        if body_items:
            for item in body_items:
                kind = item.get("type")
                value = item.get("value")
                if kind == "text" and value and str(value).strip():
                    _body(doc, str(value).strip())
                elif kind == "image" and value is not None:
                    blob = getattr(value, "blob", None)
                    if blob:
                        _add_image_to_doc(doc, blob, getattr(value, "content_type", ""),
                                          caption=getattr(value, "caption", "") or "")
                elif kind == "table" and value is not None:
                    rows = getattr(value, "rows", None) or []
                    if rows:
                        _table(doc, rows)
        else:
            for p_text in sec.get("paragraphs", []):
                if p_text.strip():
                    _body(doc, p_text.strip())
            for tbl_data in sec.get("tables", []):
                if isinstance(tbl_data, list) and tbl_data:
                    _table(doc, tbl_data)
            matched_images = image_map.get(sec_heading, [])
            for img in matched_images:
                _add_image_to_doc(doc, img.blob, img.content_type,
                                  caption=getattr(img, "caption", "") or "")

        children = sec.get("children", [])
        if children:
            # Share counters — children at deeper levels should continue
            # from the parent's counter state
            _render_appendix_sections(
                doc, children, body_headings, image_map, counters,
                apply_content_boundary=False,
            )


# ====================================================================
# Table detection helpers
# ====================================================================

def _is_table_row(line: str) -> bool:
    """Check if a line looks like a table row (pipe-separated or tab-separated)."""
    return '|' in line and line.count('|') >= 2


_SEPARATOR_CELL_RE = re.compile(r'^[\s\-:]+$')


def _is_separator_row(cells: List[str]) -> bool:
    """Detect markdown table separator rows like |---|---|, |:--|--:|."""
    if not cells:
        return False
    return all(_SEPARATOR_CELL_RE.match(c or "") for c in cells)


def _parse_table_lines(lines: List[str]) -> Optional[List[List[str]]]:
    """Parse pipe-separated or tab-separated lines into a table.

    Skips markdown separator rows (|---|---|) so they don't appear as data
    rows in the rendered Word table.
    """
    if not lines or not _is_table_row(lines[0]):
        return None
    rows = []
    for line in lines:
        if _is_table_row(line):
            cells = [c.strip() for c in line.split('|') if c.strip() or True]
            # Strip the leading/trailing empty cells introduced by | at row edges
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            if not cells:
                continue
            if _is_separator_row(cells):
                continue
            rows.append(cells)
        elif not line.strip():
            break
        else:
            break
    return rows if len(rows) >= 2 else None


# ====================================================================
# Image matching helpers
# ====================================================================

def _build_image_map(all_images_by_doc, source_filename: str) -> dict:
    """Build a dict mapping section_heading -> list of Image objects."""
    image_map = {}
    doc_images = all_images_by_doc.get(source_filename, [])
    for img in doc_images:
        if hasattr(img, 'section_heading') and img.section_heading:
            key = img.section_heading.strip()
            if key not in image_map:
                image_map[key] = []
            image_map[key].append(img)
    return image_map


def _match_body_images(all_images_by_doc, chapter_heading: str) -> list:
    """Find images from all source docs whose section_heading matches chapter_heading."""
    from doc_parser import deduplicate_images
    matched = []
    ch_clean = chapter_heading.strip()
    for doc_name, img_list in all_images_by_doc.items():
        for img in img_list:
            if hasattr(img, 'section_heading') and img.section_heading:
                sh = img.section_heading.strip()
                if sh and (sh in ch_clean or ch_clean in sh):
                    matched.append(img)
    if matched:
        groups = deduplicate_images(matched, threshold=5)
        return [group[0] for group in groups]
    return []


# ====================================================================
# Main generator
# ====================================================================

def generate_merged_docx(merge_plan, docs_data,
                         all_images_by_doc, output_path,
                         cover_title="",
                         skeleton=None,
                         template_path=None) -> str:

    # ================================================================
    # 0. START WITH TEMPLATE (if available) — preserves cover 100%
    # ================================================================
    if template_path and os.path.exists(template_path):
        shutil.copy(template_path, output_path)
        doc = Document(output_path)
        _ensure_styles(doc)
        _detach_heading_numbering(doc)
        _trim_after_cover(doc)
        final_title = merge_plan.cover_title or cover_title or "操作规程"
        _replace_cover_title(doc, final_title)
    else:
        doc = Document()
        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        try:
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception:
            pass
        if skeleton and skeleton.page_layout and skeleton.page_layout.margin_top:
            for s in doc.sections:
                s.top_margin = skeleton.page_layout.margin_top
                s.bottom_margin = skeleton.page_layout.margin_bottom
                s.left_margin = skeleton.page_layout.margin_left
                s.right_margin = skeleton.page_layout.margin_right
        else:
            for s in doc.sections:
                s.top_margin = Cm(2.5)
                s.bottom_margin = Cm(2.5)
                s.left_margin = Cm(2.8)
                s.right_margin = Cm(2.0)
        final_title = merge_plan.cover_title or cover_title or "操作规程"
        _create_default_cover(doc, final_title)

    # Build body_headings set for appendix filtering
    body_headings = set()
    for sec in merge_plan.main_sections:
        h = sec.get("heading", "")
        h_clean = re.sub(r'^[\d.、\s]+', '', h).strip()
        if h_clean:
            body_headings.add(h_clean)

    # ================================================================
    # 1. COVER (template path handled above)
    # ================================================================
    # (cover section already placed above via template copy or default cover)

    # ================================================================
    # 2. TOC — Word field code, auto-updates on open
    # ================================================================
    _insert_toc_field(doc)

    # ================================================================
    # 3. PREFACE — Heading 1, same level as body chapters
    # ================================================================
    _generate_preface(doc, template_path)

    # ================================================================
    # 4. BODY CHAPTERS — Heading 1, sub-sections Heading 2, sub-sub Heading 3
    # ================================================================
    if merge_plan.main_sections:
        chapter_idx = 0
        for section in merge_plan.main_sections:
            heading = section.get("heading", "")
            paragraphs = section.get("paragraphs", [])
            tables = section.get("tables", [])

            if not heading or _is_front_matter_heading(heading):
                continue

            clean_heading = re.sub(r'^[\d.、\s]+', '', heading).strip()
            numbered = f"{chapter_idx + 1} {clean_heading}"

            # Heading 1 for chapter title
            _h1(doc, numbered)

            # Per-chapter sub-section counters
            inner_counter = {"sub2": 0, "sub3": 0, "sub4": 0}

            # Write content with sub-section detection
            for p_text in paragraphs:
                if not p_text.strip():
                    continue

                p_lines = p_text.strip().split('\n')
                table_rows = _parse_table_lines(p_lines)
                if table_rows:
                    _table(doc, table_rows)
                    continue

                for line in p_lines:
                    line = line.strip()
                    if not line:
                        continue

                    # Detect #### / ### / ## markers for sub-sections
                    h4_marker = re.match(r'^####\s+(.+)', line)
                    h3_marker = re.match(r'^###\s+(.+)', line)
                    h2_marker = re.match(r'^##\s+(.+)', line)

                    # Detect numbered heading patterns (fallback)
                    sub2_match = re.match(r'^(\d+)\.(\d+)\s+(.+)', line)
                    sub3_match = re.match(r'^(\d+)\.(\d+)\.(\d+)\s+(.+)', line)

                    if h4_marker:
                        # If no H2/H3 exists yet, downshift to the correct level
                        if inner_counter["sub2"] == 0: inner_counter["sub2"] = 1
                        if inner_counter["sub3"] == 0: inner_counter["sub3"] = 1
                        inner_counter["sub4"] += 1
                        title = _STRIP_NUM('', h4_marker.group(1)).strip()
                        if title == clean_heading:
                            inner_counter["sub4"] -= 1
                        else:
                            new_line = f"{chapter_idx + 1}.{inner_counter['sub2']}.{inner_counter['sub3']}.{inner_counter['sub4']} {title}"
                            _h4(doc, new_line)
                    elif h3_marker:
                        # No prior H2? Treat as H2 level, not H3 (prevents "4.0.1")
                        if inner_counter["sub2"] == 0:
                            inner_counter["sub2"] += 1
                            title = _STRIP_NUM('', h3_marker.group(1)).strip()
                            if title == clean_heading:
                                inner_counter["sub2"] -= 1
                            else:
                                new_line = f"{chapter_idx + 1}.{inner_counter['sub2']} {title}"
                                _h2(doc, new_line)
                        else:
                            inner_counter["sub3"] += 1
                            inner_counter["sub4"] = 0
                            title = _STRIP_NUM('', h3_marker.group(1)).strip()
                            if title == clean_heading:
                                inner_counter["sub3"] -= 1
                            else:
                                new_line = f"{chapter_idx + 1}.{inner_counter['sub2']}.{inner_counter['sub3']} {title}"
                                _h3(doc, new_line)
                    elif h2_marker:
                        inner_counter["sub2"] += 1
                        inner_counter["sub3"] = 0
                        inner_counter["sub4"] = 0
                        title = _STRIP_NUM('', h2_marker.group(1)).strip()
                        if title == clean_heading:
                            inner_counter["sub2"] -= 1
                        else:
                            new_line = f"{chapter_idx + 1}.{inner_counter['sub2']} {title}"
                            _h2(doc, new_line)
                    elif sub3_match:
                        c, s, ss, title = sub3_match.groups()
                        inner_counter["sub2"] = int(s)
                        inner_counter["sub3"] = int(ss)
                        new_line = f"{chapter_idx + 1}.{s}.{ss} {title}"
                        _h3(doc, new_line)
                    elif sub2_match:
                        old_c, sub_n, title = sub2_match.groups()
                        inner_counter["sub2"] = int(sub_n)
                        inner_counter["sub3"] = 0
                        new_line = f"{chapter_idx + 1}.{sub_n} {title}"
                        _h2(doc, new_line)
                    elif re.match(r'^\d+[\.\、]', line):
                        _body(doc, line)
                    else:
                        _body(doc, line)

            # Dedicated tables from merge plan
            for tbl_rows in tables:
                _table(doc, tbl_rows)

            # Images matching this chapter heading
            body_images = _match_body_images(all_images_by_doc, clean_heading)
            for img in body_images:
                _add_image_to_doc(doc, img.blob, img.content_type,
                                  caption=getattr(img, "caption", "") or "")

            doc.add_paragraph()
            chapter_idx += 1

    # ================================================================
    # 5. APPENDICES
    # ================================================================
    if merge_plan.attachments:
        for att in merge_plan.attachments:
            doc.add_page_break()

            name = att.get("name", "附件")
            src_idx = att.get("source_index", -1)

            _h1(doc, name)
            doc.add_paragraph()

            # Try structured rendering from source doc sections
            if 0 <= src_idx < len(docs_data):
                src_doc = docs_data[src_idx]
                src_sections = src_doc.get("sections", [])
                src_filename = src_doc.get("filename", "")

                if src_sections:
                    image_map = _build_image_map(all_images_by_doc, src_filename)
                    _render_appendix_sections(
                        doc, src_sections, body_headings, image_map
                    )
                    continue

            # Fallback: render flat paragraphs if no structured sections
            paragraphs = att.get("paragraphs", [])
            for p_text in paragraphs:
                if p_text.strip():
                    for line in p_text.strip().split("\n"):
                        line = line.strip()
                        if line:
                            _body(doc, line)

    doc.save(output_path)
    return output_path
